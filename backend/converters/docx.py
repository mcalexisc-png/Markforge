"""DOCX converter built on python-docx.

Walks the document body in source order so paragraphs, tables, lists and
images appear exactly as authored, with bold/italic/underline and hyperlink
information preserved. Review comments and tracked changes (insertions /
deletions) are extracted as well; inserted text is kept, deleted text is
removed and counted, and both are reported in the conversion stats.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from docx.document import Document as DocxDoc
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from converters.base import BaseConverter, ConversionContext, CorruptFileError
from document_model.blocks import (
    BulletListBlock,
    CaptionBlock,
    CommentBlock,
    HeadingBlock,
    HorizontalRuleBlock,
    ImageBlock,
    NumberedListBlock,
    ParagraphBlock,
    QuoteBlock,
    TableBlock,
    TableCell,
    TableRow,
    TextRun,
)
from document_model.document import Document

_BODY_TAGS = {qn("w:p"), qn("w:tbl")}
_COMMENTS_RELTYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
_COMMENT_ANCHOR_TAGS = {qn("w:commentRangeStart"), qn("w:commentReference")}


class DocxConverter(BaseConverter):
    format = "docx"
    extensions = (".docx",)

    def __init__(self, context: ConversionContext):
        super().__init__(context)
        self._link_count = 0
        self._comments: dict[str, dict] = {}
        self._emitted_comments: set[str] = set()
        self._insertions = 0
        self._deleted_chars = 0

    def convert(self) -> Document:
        path = Path(self.context.source_path)
        try:
            d = DocxDocument(str(path))
        except Exception as exc:
            raise CorruptFileError("The DOCX file could not be opened or is corrupted.") from exc
        return self._convert(d)

    def _convert(self, doc: DocxDoc) -> Document:
        props = doc.core_properties
        metadata = {
            "title": props.title,
            "author": props.author,
            "subject": props.subject,
            "keywords": props.keywords,
            "comments": props.comments,
            "created": props.created,
            "modified": props.modified,
            "last_modified_by": props.last_modified_by,
        }
        doc_model = self._build_document(metadata)
        blocks: list = []
        stats = doc_model.stats

        self._load_comments(doc)

        body = doc.element.body
        total_children = len(body)
        processed = 0
        for child in body:
            if child.tag not in _BODY_TAGS:
                continue
            processed += 1
            self.context.progress("extract", processed, total_children, "Extracting document contents")

            if child.tag == qn("w:tbl"):
                self._convert_table(DocxTable(child, doc), blocks, doc_model)
            else:
                self._convert_paragraph(DocxParagraph(child, doc), blocks, doc_model)

        if self._insertions or self._deleted_chars:
            doc_model.warnings.append(
                {
                    "code": "tracked_changes",
                    "message": (
                        f"Tracked changes found: {self._insertions} insertion(s) kept, "
                        f"{self._deleted_chars} deleted character(s) removed."
                    ),
                    "severity": "info",
                }
            )

        stats.links = self._link_count
        doc_model.blocks = blocks
        doc_model.stats = stats
        return doc_model

    def _load_comments(self, doc: DocxDoc) -> None:
        """Parse word/comments.xml (when present) into {id: {author, date, text}}."""
        from xml.etree import ElementTree as ET

        part = doc.part
        comments_part = None
        for rel in part.rels.values():
            if rel.reltype == _COMMENTS_RELTYPE:
                comments_part = rel.target_part
                break
        if comments_part is None:
            return
        try:
            root = ET.fromstring(comments_part.blob)
        except Exception:
            return
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        for comment_el in root.findall("w:comment", ns):
            cid = comment_el.get(qn("w:id"))
            if cid is None:
                continue
            author = comment_el.get(qn("w:author")) or ""
            date = comment_el.get(qn("w:date")) or ""
            paragraphs = [p for p in comment_el.findall("w:p", ns) if p.find("w:r", ns) is not None]
            texts = []
            for p in paragraphs:
                text = "".join(
                    t.text or ""
                    for t in p.findall(".//w:t", ns)
                    if t.text and t.text.strip()
                )
                if text.strip():
                    texts.append(text.strip())
            self._comments[cid] = {
                "author": author,
                "date": date,
                "text": "\n".join(texts),
            }

    # ------------------------------------------------------------------ #

    def _convert_paragraph(
        self,
        paragraph: DocxParagraph,
        blocks: list,
        doc_model: Document,
    ) -> None:
        stats = doc_model.stats
        text = paragraph.text.strip()
        style_name = (paragraph.style.name if paragraph.style else "") or ""
        normalized_style = style_name.lower()

        if normalized_style.startswith(("heading", "title")):
            level = 1
            if normalized_style.startswith("heading"):
                try:
                    level = int(style_name.replace("Heading", "").strip() or "1")
                except ValueError:
                    level = 1
            level = max(1, min(6, level))
            runs = self._extract_runs(paragraph)
            if runs:
                blocks.append(HeadingBlock(level=level, content=runs))
                stats.headings += 1
            self._attach_comments(paragraph, blocks, doc_model)
            return

        if normalized_style.startswith(("quote", "blockquote")):
            runs = self._extract_runs(paragraph)
            if runs:
                blocks.append(QuoteBlock(content=runs))
                stats.paragraphs += 1
            self._attach_comments(paragraph, blocks, doc_model)
            return

        if normalized_style.startswith("caption"):
            runs = self._extract_runs(paragraph)
            if runs:
                blocks.append(CaptionBlock(content=runs))
            self._attach_comments(paragraph, blocks, doc_model)
            return

        if text == "---" or text == "***":
            blocks.append(HorizontalRuleBlock())
            self._attach_comments(paragraph, blocks, doc_model)
            return

        num_pr = None
        ppr = paragraph._p.find(qn("w:pPr"))
        if ppr is not None:
            num_pr = ppr.find(qn("w:numPr"))
        list_kind = None
        if num_pr is not None or normalized_style.startswith("list bullet"):
            list_kind = "bullet"
        elif normalized_style.startswith("list number"):
            list_kind = "numbered"
        if list_kind == "bullet":
            ilvl_el = num_pr.find(qn("w:ilvl")) if num_pr is not None else None
            level = int(ilvl_el.get(qn("w:val"), "0")) if ilvl_el is not None else 0
            runs = self._extract_runs(paragraph)
            if not runs:
                return
            prefix = "\t" * min(level, 4)
            item = [TextRun(text=prefix), *runs] if prefix else runs
            last = blocks[-1] if blocks else None
            if isinstance(last, BulletListBlock) and last.metadata.get("level") == level:
                last.items.append(item)
            else:
                block = BulletListBlock(items=[item])
                block.metadata["level"] = level
                blocks.append(block)
                stats.lists += 1
            stats.paragraphs += 1
            self._attach_comments(paragraph, blocks, doc_model)
            return
        if list_kind == "numbered":
            runs = self._extract_runs(paragraph)
            if runs:
                blocks.append(NumberedListBlock(items=[runs]))
                stats.lists += 1
                stats.paragraphs += 1
            self._attach_comments(paragraph, blocks, doc_model)
            return

        runs = self._extract_runs(paragraph)
        if runs:
            blocks.append(ParagraphBlock(content=runs))
            stats.paragraphs += 1

        self._attach_comments(paragraph, blocks, doc_model)
        self._extract_paragraph_images(paragraph, blocks, doc_model)

    def _attach_comments(self, paragraph: DocxParagraph, blocks: list, doc_model: Document) -> None:
        """Append CommentBlocks anchored to this paragraph (each comment once)."""
        stats = doc_model.stats
        if not self._comments:
            return
        anchored: set[str] = set()
        for el in paragraph._p.iter():
            if el.tag in _COMMENT_ANCHOR_TAGS:
                cid = el.get(qn("w:id"))
                if cid is not None:
                    anchored.add(cid)
        for cid in anchored:
            if cid in self._emitted_comments:
                continue
            comment = self._comments.get(cid)
            if not comment or not comment["text"]:
                continue
            self._emitted_comments.add(cid)
            content = [TextRun(text=line.strip()) for line in comment["text"].splitlines() if line.strip()]
            if content:
                blocks.append(
                    CommentBlock(
                        content=content,
                        author=comment["author"],
                        date=comment["date"],
                    )
                )
                stats.comments += 1

    def _extract_runs(self, paragraph: DocxParagraph) -> list[TextRun]:
        """Runs in order, resolving hyperlink targets and nested containers."""
        runs: list[TextRun] = []
        for child in paragraph._p.iterchildren():
            runs.extend(self._container_runs(child, paragraph.part))
        return runs

    def _container_runs(self, el, part) -> list[TextRun]:
        """Recursively collect runs from w:r, w:hyperlink, w:ins, w:smartTag
        and field containers; w:del content is counted and dropped."""
        tag = el.tag
        if tag == qn("w:r"):
            return self._run_to_runs(el)
        if tag == qn("w:hyperlink"):
            href = ""
            rel_id = el.get(qn("r:id"))
            if rel_id:
                rel = part.rels.get(rel_id)
                href = rel.target_ref if rel else ""
                self._link_count += 1
            runs: list[TextRun] = []
            for child in el:
                if child.tag == qn("w:r"):
                    runs.extend(self._run_to_runs(child, href=href))
            return runs
        if tag == qn("w:ins"):
            self._insertions += 1
            runs = []
            for child in el:
                if child.tag in (qn("w:r"), qn("w:hyperlink"), qn("w:smartTag")):
                    runs.extend(self._container_runs(child, part))
            return runs
        if tag == qn("w:del"):
            self._deleted_chars += len("".join(t.text or "" for t in el.iter(qn("w:delText"))))
            return []
        if tag in (qn("w:smartTag"), qn("w:fldSimple")):
            runs = []
            for child in el:
                runs.extend(self._container_runs(child, part))
            return runs
        return []

    def _run_to_runs(self, run_el, href: str | None = None) -> list[TextRun]:
        """A single w:r can contain text, tabs and breaks."""
        text_parts: list[str] = []
        for child in run_el:
            if child.tag == qn("w:t"):
                text_parts.append(child.text or "")
            elif child.tag == qn("w:tab"):
                text_parts.append("\t")
            elif child.tag == qn("w:br"):
                if child.get(qn("w:type")) == "page":
                    text_parts.append("\n\n---\n\n")
                else:
                    text_parts.append("\n")
        if not text_parts:
            return []
        rpr = run_el.find(qn("w:rPr"))
        bold = italic = underline = strikethrough = code = False
        if rpr is not None:
            bold = rpr.find(qn("w:b")) is not None
            italic = rpr.find(qn("w:i")) is not None
            underline = rpr.find(qn("w:u")) is not None
            strikethrough = rpr.find(qn("w:strike")) is not None
            vert = rpr.find(qn("w:vertAlign"))
            code = vert is not None and vert.get(qn("w:val")) == "superscript"
        return [
            TextRun(
                text="".join(text_parts),
                bold=bold,
                italic=italic,
                underline=underline,
                strikethrough=strikethrough,
                code=code,
                href=href,
            )
        ]

    def _extract_paragraph_images(self, paragraph: DocxParagraph, blocks: list, doc_model: Document) -> None:
        for blip in paragraph._p.findall(".//" + qn("a:blip")):
            self._save_blip(blip, paragraph.part, blocks, doc_model)

    def _save_blip(self, blip, part, blocks: list, doc_model: Document) -> None:
        rel_id = blip.get(qn("r:embed"))
        if not rel_id:
            return
        try:
            image_part = part.related_parts.get(rel_id)
            if image_part is None:
                return
            blob = image_part.blob
        except KeyError:
            return
        if not blob:
            return
        ext = Path(image_part.partname.filename).suffix.lstrip(".")
        rel = self.context.save_image(blob, ext)
        if rel:
            doc_model.stats.images += 1
            blocks.append(ImageBlock(path=rel, alt=""))

    # ------------------------------------------------------------------ #

    def _convert_table(self, table: DocxTable, blocks: list, doc_model: Document) -> None:
        stats = doc_model.stats
        grid_span = max((len(row.cells) for row in table.rows), default=0)
        rows: list[TableRow] = []
        for row in table.rows:
            cells: list[TableCell] = []
            for cell in row.cells:
                text_runs: list[TextRun] = []
                paragraphs = [p for p in cell.paragraphs if p.text.strip()]
                for index, paragraph in enumerate(paragraphs):
                    for run in paragraph.runs:
                        if run.text:
                            text_runs.append(
                                TextRun(
                                    text=run.text,
                                    bold=bool(run.bold),
                                    italic=bool(run.italic),
                                    underline=bool(run.underline),
                                )
                            )
                    if index < len(paragraphs) - 1:
                        text_runs.append(TextRun(text="  \n"))
                cells.append(TableCell(content=text_runs))
            while len(cells) < grid_span:
                cells.append(TableCell(content=[]))
            rows.append(TableRow(cells=cells))

        if rows:
            if self.context.settings.convert_tables:
                blocks.append(TableBlock(rows=rows, has_header=False))
                stats.tables += 1
            else:
                for row in rows:
                    values = ["".join(r.text for r in cell.content) for cell in row.cells]
                    stats.paragraphs += 1
                    blocks.append(ParagraphBlock(content=[TextRun(text=" | ".join(values))]))

        for row in table.rows:
            for cell in row.cells:
                for blip in cell._tc.findall(".//" + qn("a:blip")):
                    self._save_blip(blip, table.part, blocks, doc_model)
